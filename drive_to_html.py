#!/usr/bin/env python3
"""
drive_to_html.py  ·  Nexa CRM — Erdoğan Işık Özel
────────────────────────────────────────────────────────────
Google Drive klasöründeki TÜM dosyaları okuyup
lüks bir HTML sunum sayfası üretir.

Desteklenen dosya türleri:
  • Görseller   : jpg, jpeg, png, gif, webp, svg
  • PDF         : metin + sayfa sayısı
  • Word        : .docx paragraf / başlıklar
  • Excel/CSV   : tablo verisi
  • JSON        : yapılandırılmış veri
  • Diğerleri   : metadata kartı

Auth seçenekleri (öncelik sırası):
  1. GOOGLE_SERVICE_ACCOUNT  env var  → JSON string (önerilen, Render için)
  2. GOOGLE_SERVICE_ACCOUNT_FILE env → dosya yolu
  3. credentials.json dosyası        → OAuth akışı (lokal geliştirme)

Gerekli env variable'lar:
  GDRIVE_FOLDER_ID          → Google Drive klasör ID'si
  GOOGLE_SERVICE_ACCOUNT    → Service account JSON string
  PROJE_ADI                 → (opsiyonel) Başlık override
  PROJE_ALT_BASLIK          → (opsiyonel) Alt başlık override

Kurulum:
  pip install google-api-python-client google-auth-httplib2 \
              google-auth-oauthlib pandas openpyxl python-docx \
              PyPDF2 Pillow
────────────────────────────────────────────────────────────
"""

import os, io, json, base64, textwrap, html, re, hashlib, pickle, shutil
from pathlib import Path
from datetime import datetime

# ── Google API ──────────────────────────────────────────────
from google.oauth2.credentials import Credentials
from google.oauth2 import service_account as _sa
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload

# ── Dosya işleme ────────────────────────────────────────────
import pandas as pd
from docx import Document as DocxDocument
import PyPDF2
from PIL import Image

# ════════════════════════════════════════════════════════════
#  AYARLAR  —  Tümü env variable ile override edilebilir
# ════════════════════════════════════════════════════════════
FOLDER_ID        = os.environ.get("GDRIVE_FOLDER_ID",   "1wl6IORLksewhrWqpCOfjFNgjlC_rAhZT")
SCOPES           = ["https://www.googleapis.com/auth/drive.readonly"]
TOKEN_FILE       = "token.json"
CREDS_FILE       = "credentials.json"
OUTPUT_FILE      = "sunum.html"
PROJE_ADI        = os.environ.get("PROJE_ADI",        "Erdoğan Işık Özel · Proje Portföyü")
PROJE_ALT_BASLIK = os.environ.get("PROJE_ALT_BASLIK", "CB VIP Dikmen — Premium Gayrimenkul Arşivi")
CACHE_DIR        = ".drive_cache"
ASSETS_DIR       = "sunum_assets"


# ════════════════════════════════════════════════════════════
#  1. GOOGLE DRIVE KİMLİK DOĞRULAMA
#  Öncelik: Service Account env var → SA dosya → OAuth akışı
# ════════════════════════════════════════════════════════════
def get_service():
    """
    Google Drive API bağlantısı kurar.
    Öncelik sırası:
      1. GOOGLE_SERVICE_ACCOUNT env var → Service Account JSON string (en güvenli)
      2. GOOGLE_OAUTH_TOKEN env var     → token.json içeriği (OAuth, refresh_token ile yenilenir)
      3. token.json dosyası             → lokal geliştirme
      4. credentials.json + OAuth flow  → lokal ilk kurulum
    """
    # ── 1. Service Account — JSON string (Render env var) ────
    sa_json_str = os.environ.get("GOOGLE_SERVICE_ACCOUNT", "").strip()
    if sa_json_str.startswith("{"):
        try:
            sa_info = json.loads(sa_json_str)
            creds   = _sa.Credentials.from_service_account_info(sa_info, scopes=SCOPES)
            print("  ✓ Auth: Service Account (env JSON)")
            return build("drive", "v3", credentials=creds)
        except Exception as e:
            print(f"  ⚠  SA env parse hatası: {e} — OAuth'a geçiliyor…")

    # ── 2. Service Account — dosya yolu ─────────────────────
    sa_file = os.environ.get("GOOGLE_SERVICE_ACCOUNT_FILE", "").strip()
    if sa_file and Path(sa_file).exists():
        try:
            creds = _sa.Credentials.from_service_account_file(sa_file, scopes=SCOPES)
            print(f"  ✓ Auth: Service Account (dosya: {sa_file})")
            return build("drive", "v3", credentials=creds)
        except Exception as e:
            print(f"  ⚠  SA dosya hatası: {e} — OAuth'a geçiliyor…")

    # ── 3. GOOGLE_OAUTH_TOKEN env var → token.json'ı diske yaz ──
    oauth_token_str = os.environ.get("GOOGLE_OAUTH_TOKEN", "").strip()
    if oauth_token_str.startswith("{"):
        try:
            token_path = Path(TOKEN_FILE)
            token_path.write_text(oauth_token_str, encoding="utf-8")
            print("  ✓ GOOGLE_OAUTH_TOKEN env var → token.json'a yazıldı")
        except Exception as e:
            print(f"  ⚠  GOOGLE_OAUTH_TOKEN yazma hatası: {e}")

    # ── 4. OAuth akışı (token.json veya credentials.json) ───
    creds = None
    if Path(TOKEN_FILE).exists():
        try:
            creds = Credentials.from_authorized_user_file(TOKEN_FILE, SCOPES)
        except Exception as e:
            print(f"  ⚠  token.json okunamadı: {e}")

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            print("  🔄 Token süresi dolmuş, yenileniyor…")
            creds.refresh(Request())
            # Yenilenen token'ı diske kaydet (Render'da ephemeral ama process boyunca işe yarar)
            try:
                with open(TOKEN_FILE, "w") as f:
                    f.write(creds.to_json())
                print("  ✓ Token yenilendi ve token.json güncellendi")
            except Exception as e:
                print(f"  ⚠  token.json güncellenemedi: {e}")
        else:
            if not Path(CREDS_FILE).exists():
                raise FileNotFoundError(
                    f"Google Drive kimlik doğrulaması başarısız!\n"
                    f"Seçenekler:\n"
                    f"  A) Render → Environment → GOOGLE_OAUTH_TOKEN = token.json içeriği (JSON string)\n"
                    f"  B) Render → Environment → GOOGLE_SERVICE_ACCOUNT = service account JSON\n"
                    f"  C) Lokal: credentials.json dosyasını proje köküne koy\n"
                    f"token.json mevcut: {Path(TOKEN_FILE).exists()} | "
                    f"credentials.json mevcut: {Path(CREDS_FILE).exists()}"
                )
            flow  = InstalledAppFlow.from_client_secrets_file(CREDS_FILE, SCOPES)
            creds = flow.run_local_server(port=0)
            with open(TOKEN_FILE, "w") as f:
                f.write(creds.to_json())

    print("  ✓ Auth: OAuth (token.json)")
    return build("drive", "v3", credentials=creds)


# ════════════════════════════════════════════════════════════
#  2. DOSYA LİSTESİ ÇEK (özyinelemeli – alt klasörlere iner)
# ════════════════════════════════════════════════════════════
def list_files(service, folder_id, _depth=0, _path=""):
    """Verilen klasördeki tüm dosyaları özyinelemeli olarak döner.
    Alt klasörler atlanır; içindeki dosyalar 'folder_path' alanıyla döner."""
    results, page_token = [], None
    while True:
        resp = service.files().list(
            q=f"'{folder_id}' in parents and trashed=false",
            spaces="drive",
            fields="nextPageToken, files(id,name,mimeType,size,modifiedTime)",
            pageToken=page_token
        ).execute()
        items = resp.get("files", [])
        for item in items:
            if item.get("mimeType") == "application/vnd.google-apps.folder":
                # Alt klasörün içine in
                sub_path = (_path + " / " if _path else "") + item["name"]
                print(f"  {'  ' * _depth}📁 {item['name']} (alt klasör, içine giriliyor…)")
                sub_files = list_files(service, item["id"], _depth + 1, sub_path)
                results.extend(sub_files)
            else:
                item["folder_path"] = _path  # hangi klasörden geldiğini tut
                results.append(item)
        page_token = resp.get("nextPageToken")
        if not page_token:
            break
    return results


# ════════════════════════════════════════════════════════════
#  3. DOSYA İNDİR
#  Google Workspace dosyaları (Slides/Docs/Sheets) export gerektirir
# ════════════════════════════════════════════════════════════

# Google Workspace MIME → export MIME + uzantı
GAPPS_EXPORT = {
    "application/vnd.google-apps.presentation": ("application/pdf", "pdf"),
    "application/vnd.google-apps.document":     ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"),
    "application/vnd.google-apps.spreadsheet":  ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "xlsx"),
    "application/vnd.google-apps.drawing":      ("image/png", "png"),
}

def download_bytes(service, file_id, mime_type="") -> tuple[bytes, str]:
    """(bytes, gerçek_mime) döner. Google Workspace dosyaları export edilir.
    Mime bilinmiyorsa veya 403 alınırsa otomatik export formatlarını dener."""

    def _dl(req) -> bytes:
        buf = io.BytesIO()
        dl = MediaIoBaseDownload(buf, req)
        done = False
        while not done:
            _, done = dl.next_chunk()
        return buf.getvalue()

    # 1. Bilinen Google Apps mime → direkt export
    if mime_type in GAPPS_EXPORT:
        export_mime, _ = GAPPS_EXPORT[mime_type]
        return _dl(service.files().export_media(fileId=file_id, mimeType=export_mime)), export_mime

    # 2. Normal binary indirmeyi dene
    try:
        return _dl(service.files().get_media(fileId=file_id)), mime_type
    except Exception as e:
        if "fileNotDownloadable" not in str(e) and "403" not in str(e):
            raise

    # 3. Fallback: mime ne olursa olsun export formatlarını sırayla dene
    for export_mime in [
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.presentationml.presentation",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "image/png",
    ]:
        try:
            data = _dl(service.files().export_media(fileId=file_id, mimeType=export_mime))
            return data, export_mime
        except Exception:
            continue

    raise RuntimeError(f"Dosya indirilemedi: file_id={file_id}, mime={mime_type}")


# ════════════════════════════════════════════════════════════
#  3b. ÖNBELLEK  (Cache)
# ════════════════════════════════════════════════════════════
def _cache_key(file_id: str, modified_time: str) -> str:
    """file_id + modifiedTime → benzersiz anahtar (hex)"""
    raw = f"{file_id}_{modified_time}".encode()
    return hashlib.sha1(raw).hexdigest()

def cache_get(file_id: str, modified_time: str):
    """Cache'de varsa (bytes, mime) döner, yoksa None."""
    key  = _cache_key(file_id, modified_time)
    path = Path(CACHE_DIR) / key
    if path.exists():
        try:
            with open(path, "rb") as f:
                return pickle.load(f)   # (bytes, mime_str)
        except Exception:
            path.unlink(missing_ok=True)
    return None

def cache_set(file_id: str, modified_time: str, data: bytes, mime: str):
    """(bytes, mime) önbelleğe yazar."""
    Path(CACHE_DIR).mkdir(exist_ok=True)
    key  = _cache_key(file_id, modified_time)
    path = Path(CACHE_DIR) / key
    with open(path, "wb") as f:
        pickle.dump((data, mime), f)

def cache_stats() -> dict:
    """Önbellek klasörü hakkında istatistik."""
    d = Path(CACHE_DIR)
    if not d.exists():
        return {"files": 0, "size_mb": 0.0}
    files = list(d.iterdir())
    total = sum(p.stat().st_size for p in files if p.is_file())
    return {"files": len(files), "size_mb": total / (1024*1024)}

def download_cached(service, file_id: str, mime_type: str, modified_time: str) -> tuple[bytes, str]:
    """Önce cache bakar; yoksa Drive'dan indirir ve önbelleğe yazar."""
    hit = cache_get(file_id, modified_time)
    if hit:
        return hit          # (bytes, mime)
    data, real_mime = download_bytes(service, file_id, mime_type)
    cache_set(file_id, modified_time, data, real_mime)
    return data, real_mime


# ════════════════════════════════════════════════════════════
#  4. DOSYA TİPİNE GÖRE İŞLE → HTML KART
# ════════════════════════════════════════════════════════════
def ext(name: str) -> str:
    return Path(name).suffix.lower().lstrip(".")

def size_fmt(b) -> str:
    if b is None: return "—"
    b = int(b)
    for unit in ("B","KB","MB","GB"):
        if b < 1024: return f"{b:.1f} {unit}"
        b /= 1024
    return f"{b:.1f} TB"

# ── Görsel ──────────────────────────────────────────────────
def process_image(data: bytes, name: str, mime: str) -> str:
    # thumbnail oluştur
    try:
        img = Image.open(io.BytesIO(data))
        img.thumbnail((900, 600))
        buf = io.BytesIO()
        fmt = "PNG" if name.lower().endswith(".png") else "JPEG"
        img.save(buf, format=fmt)
        b64 = base64.b64encode(buf.getvalue()).decode()
        src = f"data:{mime};base64,{b64}"
    except Exception:
        b64 = base64.b64encode(data).decode()
        src = f"data:{mime};base64,{b64}"
    return f'<div class="card card-image"><div class="card-label">🖼 Görsel</div>' \
           f'<h3 class="card-title">{html.escape(name)}</h3>' \
           f'<img src="{src}" alt="{html.escape(name)}" loading="lazy"/></div>'

# ── PDF (iframe embed — scrollable, open by default) ────────
def process_pdf(data: bytes, name: str, label: str = "PDF") -> str:
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(data))
        pages  = len(reader.pages)
    except Exception:
        pages = 0
    uid     = hashlib.md5(data[:128]).hexdigest()[:12]
    pdf_path = save_pdf(data, uid)
    lbl     = "🎞 Sunum" if label == "Slides" else "📄 PDF"
    pg_tag  = f"<span class='pdf-pg'>{pages} sayfa</span>" if pages else ""
    dl_attr = f'href="{pdf_path}" download="{html.escape(name)}"'
    return f"""<div class="card card-pdf-embed" id="card-{uid}">
  <div class="pdf-embed-header">
    <span class="card-label">{lbl}</span>
    {pg_tag}
    <span class="pdf-title">{html.escape(name)}</span>
    <div class="pdf-embed-actions">
      <button class="pdf-btn" onclick="pdfPrev('{uid}')">‹ Önceki</button>
      <span class="pdf-page-info" id="pginfo-{uid}">1 / {pages if pages else '?'}</span>
      <button class="pdf-btn" onclick="pdfNext('{uid}')">Sonraki ›</button>
      <button class="pdf-btn" onclick="togglePdfFull('{uid}')">⤢ Tam Ekran</button>
      <a {dl_attr} class="pdf-btn">⬇ İndir</a>
    </div>
  </div>
  <div class="pdf-embed-wrap" id="pdfwrap-{uid}">
    <div class="pdf-canvas-wrap" id="cvwrap-{uid}">
      <canvas id="cv-{uid}" class="pdf-canvas"></canvas>
      <div class="pdf-loading" id="pdfload-{uid}">PDF yükleniyor…</div>
    </div>
  </div>
</div>
<script>
(function(){{
  const PDF_URL="{pdf_path}";
  const UID="{uid}";
  const TOTAL={pages if pages else 0};
  let pdfDoc=null, curPage=1, rendering=false, pendingPage=null;

  function renderPage(num){{
    if(rendering){{ pendingPage=num; return; }}
    rendering=true;
    pdfDoc.getPage(num).then(function(page){{
      const wrap=document.getElementById('cvwrap-'+UID);
      const scale=Math.min((wrap.clientWidth||900)/page.getViewport({{scale:1}}).width, 2.5);
      const vp=page.getViewport({{scale:scale}});
      const canvas=document.getElementById('cv-'+UID);
      canvas.height=vp.height; canvas.width=vp.width;
      page.render({{canvasContext:canvas.getContext('2d'), viewport:vp}}).promise.then(function(){{
        rendering=false;
        document.getElementById('pginfo-'+UID).textContent=num+' / '+(TOTAL||pdfDoc.numPages);
        document.getElementById('pdfload-'+UID).style.display='none';
        if(pendingPage){{ const p=pendingPage; pendingPage=null; renderPage(p); }}
      }});
    }});
  }}

  function init(){{
    if(!window.pdfjsLib){{ setTimeout(init,120); return; }}
    window.pdfjsLib.GlobalWorkerOptions.workerSrc=
      'https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.worker.min.js';
    window.pdfjsLib.getDocument(PDF_URL).promise.then(function(doc){{
      pdfDoc=doc; curPage=1; renderPage(1);
    }}).catch(function(e){{
      document.getElementById('pdfload-'+UID).textContent='PDF açılamadı: '+e.message;
    }});
  }}

  window['_pdfState_'+UID]={{
    prev:function(){{ if(!pdfDoc||curPage<=1) return; curPage--; renderPage(curPage); }},
    next:function(){{ if(!pdfDoc||curPage>=(TOTAL||pdfDoc.numPages)) return; curPage++; renderPage(curPage); }}
  }};

  if(document.readyState==='loading') document.addEventListener('DOMContentLoaded',init);
  else init();
}})();
</script>"""

# ── Word ─────────────────────────────────────────────────────
def process_docx(data: bytes, name: str) -> str:
    try:
        doc  = DocxDocument(io.BytesIO(data))
        rows = []
        for para in doc.paragraphs[:40]:
            t = para.text.strip()
            if not t: continue
            style = para.style.name.lower()
            if "heading 1" in style:
                rows.append(f"<h4 class='doc-h1'>{html.escape(t)}</h4>")
            elif "heading 2" in style:
                rows.append(f"<h5 class='doc-h2'>{html.escape(t)}</h5>")
            elif "heading 3" in style:
                rows.append(f"<h5 class='doc-h3'>{html.escape(t)}</h5>")
            else:
                rows.append(f"<p class='doc-p'>{html.escape(t[:300])}{'…' if len(t)>300 else ''}</p>")
        body = "\n".join(rows) or "<p class='muted'>İçerik bulunamadı.</p>"
    except Exception as e:
        body = f"<p class='muted'>Hata: {e}</p>"
    return (f'<div class="card card-word"><div class="card-label">📝 Word</div>'
            f'<h3 class="card-title">{html.escape(name)}</h3>'
            f'<div class="doc-body">{body}</div></div>')

# ── Excel / CSV ──────────────────────────────────────────────
def process_table(data: bytes, name: str) -> str:
    try:
        e = ext(name)
        if e == "csv":
            df = pd.read_csv(io.BytesIO(data), nrows=50)
        else:
            df = pd.read_excel(io.BytesIO(data), nrows=50)
        rows_total = len(df)
        df = df.head(20)
        thead = "<tr>" + "".join(f"<th>{html.escape(str(c))}</th>" for c in df.columns) + "</tr>"
        tbody = ""
        for _, row in df.iterrows():
            tbody += "<tr>" + "".join(f"<td>{html.escape(str(v))}</td>" for v in row) + "</tr>"
        note = f"<p class='table-note'>İlk 20 satır gösteriliyor · Toplam: {rows_total}</p>" if rows_total>20 else ""
    except Exception as e:
        thead, tbody, note = "", "", f"<p class='muted'>Hata: {e}</p>"
    return (f'<div class="card card-table"><div class="card-label">📊 {"Excel" if ext(name)!="csv" else "CSV"}</div>'
            f'<h3 class="card-title">{html.escape(name)}</h3>'
            f'<div class="table-wrap"><table><thead>{thead}</thead><tbody>{tbody}</tbody></table></div>'
            f'{note}</div>')

# ── JSON ─────────────────────────────────────────────────────
def process_json(data: bytes, name: str) -> str:
    try:
        obj = json.loads(data.decode("utf-8", errors="replace"))
        pretty = json.dumps(obj, ensure_ascii=False, indent=2)
        snippet = pretty[:1200] + ("\n…" if len(pretty) > 1200 else "")
    except Exception as e:
        snippet = f"Hata: {e}"
    return (f'<div class="card card-json"><div class="card-label">⚙ JSON</div>'
            f'<h3 class="card-title">{html.escape(name)}</h3>'
            f'<pre class="json-pre">{html.escape(snippet)}</pre></div>')

# ── Diğer ────────────────────────────────────────────────────
def process_other(meta: dict) -> str:
    folder_path = meta.get("folder_path", "")
    folder_tag = (f'<p class="meta">📁 {html.escape(folder_path)}</p>' if folder_path else "")
    return (f'<div class="card card-other"><div class="card-label">📎 Dosya</div>'
            f'<h3 class="card-title">{html.escape(meta["name"])}</h3>'
            f'{folder_tag}'
            f'<p class="meta">Boyut: {size_fmt(meta.get("size"))}</p>'
            f'<p class="meta">Tür: {html.escape(meta.get("mimeType","—"))}</p>'
            f'<p class="muted">Önizleme desteklenmiyor.</p></div>')


# ════════════════════════════════════════════════════════════
#  5. YARDIMCI
# ════════════════════════════════════════════════════════════
def slugify(s: str) -> str:
    s = s.upper()
    tr = str.maketrans("ÇĞİÖŞÜçğıöşü", "CGIOSUcgiosu")
    s = s.translate(tr)
    return re.sub(r"[^A-Z0-9]+", "_", s).strip("_")

def file_type_key(name: str, mime: str) -> str:
    e = ext(name)
    if "google-apps.presentation" in mime: return "pdf"
    if "google-apps.document"     in mime: return "word"
    if "google-apps.spreadsheet"  in mime: return "table"
    if "google-apps.drawing"      in mime: return "image"
    if "pdf"            in mime: return "pdf"
    if "presentationml" in mime: return "pdf"
    if "wordprocessing" in mime: return "word"
    if "spreadsheetml"  in mime: return "table"
    if e in ("jpg","jpeg","png","gif","webp","svg") or mime.startswith("image/"): return "image"
    if e in ("xlsx","xls","csv"): return "table"
    if e == "docx":  return "word"
    if e == "json":  return "json"
    if e == "pdf":   return "pdf"
    return "other"

# ════════════════════════════════════════════════════════════
#  6. ASSET DOSYALARI  (base64 yerine disk'e yaz)
# ════════════════════════════════════════════════════════════
def _asset_path(uid: str, ext_: str) -> Path:
    p = Path(ASSETS_DIR)
    p.mkdir(exist_ok=True)
    return p / f"{uid}.{ext_}"

def _prepare_image(data: bytes, size: tuple) -> tuple:
    """Görseli aç, boyutlandır, RGB'ye dönüştür. (img, ext) döner."""
    img = Image.open(io.BytesIO(data))
    img.thumbnail(size, Image.LANCZOS)
    # RGBA / P (palette) / LA → RGB (JPEG uyumlu)
    if img.mode in ("RGBA", "LA", "P"):
        bg = Image.new("RGB", img.size, (18, 18, 24))   # koyu arka plan
        paste_img = img.convert("RGBA") if img.mode != "RGBA" else img
        bg.paste(paste_img, mask=paste_img.split()[3])
        img = bg
    elif img.mode != "RGB":
        img = img.convert("RGB")
    return img

def make_thumb(data: bytes, mime: str, uid: str, size=(480, 320)) -> str:
    """Küçük thumbnail → sunum_assets/{uid}_t.jpg  Döner: relative path string"""
    try:
        Path(ASSETS_DIR).mkdir(exist_ok=True)
        p = Path(ASSETS_DIR) / f"{uid}_t.jpg"
        if p.exists():
            return f"{ASSETS_DIR}/{uid}_t.jpg"
        img = _prepare_image(data, size)
        img.save(str(p), format="JPEG", quality=82, optimize=True)
        return f"{ASSETS_DIR}/{uid}_t.jpg"
    except Exception as e:
        print(f"    ⚠  Thumbnail hatası ({uid}): {e}")
        return ""

def make_large(data: bytes, mime: str, uid: str, size=(1000, 750)) -> str:
    """Büyük görsel → sunum_assets/{uid}_l.jpg  Döner: relative path string"""
    try:
        Path(ASSETS_DIR).mkdir(exist_ok=True)
        p = Path(ASSETS_DIR) / f"{uid}_l.jpg"
        if p.exists():
            return f"{ASSETS_DIR}/{uid}_l.jpg"
        img = _prepare_image(data, size)
        img.save(str(p), format="JPEG", quality=88, optimize=True)
        return f"{ASSETS_DIR}/{uid}_l.jpg"
    except Exception as e:
        print(f"    ⚠  Large görsel hatası ({uid}): {e}")
        return ""

def save_video(data: bytes, uid: str) -> str:
    """Video → sunum_assets/{uid}.mp4  Döner: relative path string"""
    try:
        Path(ASSETS_DIR).mkdir(exist_ok=True)
        p = Path(ASSETS_DIR) / f"{uid}.mp4"
        p.write_bytes(data)
        return f"{ASSETS_DIR}/{uid}.mp4"
    except Exception:
        return ""

def save_pdf(data: bytes, uid: str) -> str:
    """PDF → sunum_assets/{uid}.pdf  Döner: relative path string"""
    try:
        Path(ASSETS_DIR).mkdir(exist_ok=True)
        p = Path(ASSETS_DIR) / f"{uid}.pdf"
        p.write_bytes(data)
        return f"{ASSETS_DIR}/{uid}.pdf"
    except Exception:
        return ""

# ════════════════════════════════════════════════════════════
#  8. MANIFEST  (değişiklik kontrolü)
# ════════════════════════════════════════════════════════════
MANIFEST_FILE = ".drive_manifest"

def compute_manifest(files: list) -> str:
    """Drive dosya listesinden deterministik bir hash üret.
    Her dosya için id + modifiedTime kullanılır."""
    parts = sorted(f"{f['id']}:{f.get('modifiedTime','')}" for f in files)
    return hashlib.sha256("\n".join(parts).encode()).hexdigest()

def load_manifest() -> str:
    p = Path(MANIFEST_FILE)
    return p.read_text(encoding="utf-8").strip() if p.exists() else ""

def save_manifest(h: str):
    Path(MANIFEST_FILE).write_text(h, encoding="utf-8")


# ════════════════════════════════════════════════════════════
#  9. HTML ŞABLONLARI
# ════════════════════════════════════════════════════════════
def _html_head(title: str) -> str:
    return f"""<!DOCTYPE html>
<html lang="tr">
<head>
<meta charset="UTF-8"/>
<meta name="viewport" content="width=device-width,initial-scale=1"/>
<title>{html.escape(title)}</title>
<meta name="description" content="Erdoğan Işık Özel · CB VIP Dikmen — Premium gayrimenkul proje portföyü ve sunum arşivi."/>
<link rel="icon" type="image/jpeg" href="https://i.ibb.co/Tz1MNtv/Whats-App-mage-2026-02-14-at-12-33-50.jpg"/>
<link rel="preconnect" href="https://fonts.googleapis.com"/>
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin/>
<link rel="stylesheet" href="https://fonts.googleapis.com/css2?family=Cormorant+Garamond:ital,wght@0,300;0,400;0,600;1,300;1,400;0,700&family=Raleway:wght@300;400;500;600;700&family=DM+Mono:wght@300;400&display=swap"/>
<link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.5.0/css/all.min.css"/>
<script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.min.js"></script>
<style>
"""

def _html_css() -> str:
    return """:root {
  --bg:      #080810;
  --bg2:     #0f0f1a;
  --bg3:     #141420;
  --bg4:     #1a1a28;
  --gold:    #c9a84c;
  --gold2:   #e8c97a;
  --gold3:   #f5dfa0;
  --teal:    #2dd4bf;
  --orange:  #f97316;
  --silver:  #8c8fa8;
  --text:    #ede9df;
  --text2:   #9b9aac;
  --border:  rgba(201,168,76,.15);
  --border2: rgba(201,168,76,.30);
  --r:       14px;
  --ser:     "Cormorant Garamond", Georgia, serif;
  --sans:    "Raleway", -apple-system, sans-serif;
  --mono:    "DM Mono", monospace;
  --trans:   .25s ease;
}
*,*::before,*::after{box-sizing:border-box;margin:0;padding:0}
html{scroll-behavior:smooth}
body{background:var(--bg);color:var(--text);font-family:var(--ser);font-size:17px;line-height:1.7;min-height:100vh}
body::before{content:"";position:fixed;inset:0;z-index:0;pointer-events:none;background-image:url("data:image/svg+xml,%3Csvg viewBox='0 0 200 200' xmlns='http://www.w3.org/2000/svg'%3E%3Cfilter id='n'%3E%3CfeTurbulence type='fractalNoise' baseFrequency='0.85' numOctaves='4'/%3E%3C/filter%3E%3Crect width='100%25' height='100%25' filter='url(%23n)' opacity='0.035'/%3E%3C/svg%3E");opacity:.4}
.site-nav{position:sticky;top:0;z-index:100;background:rgba(8,8,16,.92);backdrop-filter:blur(20px);border-bottom:1px solid var(--border);padding:0 48px;display:flex;align-items:center;justify-content:space-between;height:62px}
.site-nav .nav-logo{font-family:var(--ser);font-size:1.05rem;font-weight:400;color:var(--text);text-decoration:none;letter-spacing:-.01em;white-space:nowrap}
.site-nav .nav-logo em{font-style:italic;color:var(--gold2)}
.site-nav .nav-links{display:flex;align-items:center;gap:0}
.site-nav .nav-links a{font-family:var(--mono);font-size:.68rem;letter-spacing:.12em;text-transform:uppercase;color:var(--text2);text-decoration:none;padding:0 16px;height:62px;display:flex;align-items:center;border-bottom:2px solid transparent;transition:all var(--trans);white-space:nowrap}
.site-nav .nav-links a:hover{color:var(--text)}
.site-nav .nav-links a.active{color:var(--gold2);border-bottom-color:var(--gold)}
.site-nav .nav-links a.nav-wa{color:var(--teal);border:1px solid rgba(45,212,191,.25);border-radius:999px;margin-left:12px;padding:0 16px;height:34px}
.site-nav .nav-links a.nav-wa:hover{background:rgba(45,212,191,.1)}
@media(max-width:768px){.site-nav{padding:0 20px}.site-nav .nav-links a{padding:0 10px;font-size:.6rem}.site-nav .nav-links a.nav-wa{display:none}}
.site-header{position:relative;z-index:10;padding:60px 64px 48px;background:linear-gradient(155deg,#13132a 0%,#0a0a14 55%,var(--bg) 100%);border-bottom:1px solid var(--border);overflow:hidden}
.site-header::after{content:"";position:absolute;top:-60px;right:-60px;width:500px;height:500px;border-radius:50%;background:radial-gradient(circle,rgba(201,168,76,.07) 0%,transparent 70%);pointer-events:none}
.header-inner{max-width:1300px;margin:0 auto;position:relative;z-index:1}
.header-eyebrow{font-family:var(--mono);font-size:.68rem;letter-spacing:.25em;text-transform:uppercase;color:var(--gold);margin-bottom:20px}
.site-header h1{font-size:clamp(2.4rem,5vw,4.4rem);font-weight:300;letter-spacing:-.025em;line-height:1.08;color:var(--text)}
.site-header h1 em{font-style:italic;color:var(--gold2)}
.header-sub{margin-top:14px;color:var(--text2);font-size:1rem;font-weight:300;font-family:var(--sans)}
.header-meta{margin-top:32px;display:flex;gap:14px;flex-wrap:wrap;align-items:center}
.hm-pill{font-family:var(--mono);font-size:.68rem;letter-spacing:.1em;color:var(--silver);padding:5px 14px;border:1px solid var(--border);border-radius:999px;background:rgba(255,255,255,.03)}
.hm-pill span{color:var(--gold2)}
.header-refresh{margin-left:auto;font-family:var(--mono);font-size:.65rem;letter-spacing:.1em;color:var(--silver);background:transparent;border:1px solid var(--border);border-radius:999px;padding:6px 16px;cursor:pointer;transition:all var(--trans);display:inline-flex;align-items:center;gap:6px}
.header-refresh:hover{color:var(--gold);border-color:var(--gold);background:rgba(201,168,76,.07)}
.header-refresh.spinning i{animation:spin .8s linear infinite}
@keyframes spin{to{transform:rotate(360deg)}}
#homepage{position:relative;z-index:1}
.home-grid{max-width:1400px;margin:0 auto;padding:60px 48px 80px;display:grid;grid-template-columns:repeat(auto-fill,minmax(360px,1fr));gap:28px}
.proj-card{background:var(--bg2);border:1px solid var(--border);border-radius:var(--r);overflow:hidden;cursor:pointer;transition:transform var(--trans),box-shadow var(--trans),border-color var(--trans);animation:fadeUp .5s ease both}
.proj-card:hover{transform:translateY(-6px);box-shadow:0 24px 60px rgba(0,0,0,.6);border-color:var(--border2)}
.proj-card:hover .proj-open-hint{opacity:1;transform:translateY(0)}
.proj-card:focus{outline:2px solid var(--gold);outline-offset:2px}
.proj-cover-wrap{position:relative;overflow:hidden;aspect-ratio:16/9;background:var(--bg3)}
.proj-cover,.proj-cover-video{width:100%;height:100%;object-fit:cover;display:block;transition:transform .4s ease}
.proj-card:hover .proj-cover,.proj-card:hover .proj-cover-video{transform:scale(1.05)}
.proj-cover-placeholder{width:100%;height:100%;display:flex;align-items:center;justify-content:center;font-size:3rem;color:var(--border2)}
.proj-overlay{position:absolute;inset:0;background:linear-gradient(to top,rgba(8,8,16,.75) 0%,transparent 60%);display:flex;align-items:flex-end;padding:20px}
.proj-open-hint{font-family:var(--mono);font-size:.7rem;letter-spacing:.15em;color:var(--gold2);opacity:0;transform:translateY(6px);transition:all var(--trans)}
.proj-video-badge{position:absolute;top:12px;right:12px;font-family:var(--mono);font-size:.6rem;letter-spacing:.1em;color:var(--gold);background:rgba(0,0,0,.65);backdrop-filter:blur(6px);border:1px solid var(--gold);border-radius:999px;padding:4px 10px;pointer-events:none}
.proj-info{padding:24px 28px 28px}
.proj-eyebrow{font-family:var(--mono);font-size:.62rem;letter-spacing:.2em;text-transform:uppercase;color:var(--gold);margin-bottom:8px}
.proj-name{font-size:1.35rem;font-weight:400;line-height:1.25;color:var(--text);margin-bottom:14px;word-break:break-word}
.proj-badges{display:flex;flex-wrap:wrap;gap:6px;margin-bottom:12px}
.badge{font-family:var(--mono);font-size:.65rem;letter-spacing:.08em;color:var(--text2);padding:4px 10px;border:1px solid var(--border);border-radius:999px;background:rgba(255,255,255,.025)}
.badge-total{border-color:var(--gold);color:var(--gold)}
.badge-video{border-color:var(--gold2);color:var(--gold2)}
.proj-total{font-family:var(--mono);font-size:.65rem;color:var(--silver);letter-spacing:.08em}
.project-page{position:relative;z-index:1;min-height:100vh}
.page-header{padding:44px 64px 36px;background:linear-gradient(155deg,#13132a 0%,#0a0a14 55%,var(--bg) 100%);border-bottom:1px solid var(--border)}
.back-btn{font-family:var(--mono);font-size:.72rem;letter-spacing:.12em;color:var(--silver);background:transparent;border:1px solid var(--border);border-radius:999px;padding:7px 18px;cursor:pointer;transition:all var(--trans);margin-bottom:24px;display:inline-block}
.back-btn:hover{color:var(--gold);border-color:var(--gold);background:rgba(201,168,76,.08)}
.page-eyebrow{font-family:var(--mono);font-size:.65rem;letter-spacing:.22em;text-transform:uppercase;color:var(--gold);margin-bottom:10px}
.page-title{font-size:clamp(2rem,4vw,3.4rem);font-weight:300;letter-spacing:-.02em;color:var(--text);margin-bottom:16px}
.page-stats{display:flex;flex-wrap:wrap;gap:8px}
.tab-bar{position:sticky;top:62px;z-index:19;padding:0 64px;background:rgba(8,8,16,.95);backdrop-filter:blur(20px);border-bottom:1px solid var(--border);display:flex;gap:0;overflow-x:auto}
.tab-bar::-webkit-scrollbar{height:0}
.tab-btn{font-family:var(--mono);font-size:.72rem;letter-spacing:.1em;color:var(--text2);background:transparent;border:none;padding:18px 24px;cursor:pointer;border-bottom:2px solid transparent;white-space:nowrap;transition:all var(--trans)}
.tab-btn:hover{color:var(--text)}
.tab-btn.active{color:var(--gold2);border-bottom-color:var(--gold)}
.tab-count{font-size:.6rem;color:var(--silver);margin-left:5px}
.tab-content{padding:48px 64px 80px;max-width:1500px;margin:0 auto}
.gallery{display:grid;grid-template-columns:repeat(auto-fill,minmax(260px,1fr));gap:14px}
.gal-item{position:relative;overflow:hidden;border-radius:10px;background:var(--bg3);border:1px solid var(--border);cursor:zoom-in;aspect-ratio:4/3;transition:transform var(--trans),box-shadow var(--trans),border-color var(--trans)}
.gal-item:hover{transform:translateY(-3px);box-shadow:0 16px 40px rgba(0,0,0,.5);border-color:var(--border2)}
.gal-item img{width:100%;height:100%;object-fit:cover;display:block;transition:transform .35s ease}
.gal-item:hover img{transform:scale(1.04)}
.gal-caption{position:absolute;bottom:0;left:0;right:0;background:linear-gradient(transparent,rgba(0,0,0,.75));padding:24px 12px 10px;font-family:var(--mono);font-size:.6rem;letter-spacing:.06em;color:rgba(255,255,255,.75);opacity:0;transition:opacity var(--trans);word-break:break-word}
.gal-item:hover .gal-caption{opacity:1}
.gal-sub{display:block;color:var(--gold);margin-top:2px;font-size:.55rem}
.pdf-stack{display:flex;flex-direction:column;gap:32px}
.card-pdf-embed{background:var(--bg2);border:1px solid var(--border);border-radius:var(--r);overflow:hidden;transition:box-shadow var(--trans),border-color var(--trans)}
.card-pdf-embed:hover{box-shadow:0 16px 40px rgba(0,0,0,.45);border-color:var(--border2)}
.pdf-embed-header{display:flex;align-items:center;gap:10px;flex-wrap:wrap;padding:14px 20px;background:var(--bg3);border-bottom:1px solid var(--border)}
.pdf-embed-header .card-label{margin-bottom:0;font-size:.62rem}
.pdf-pg{font-family:var(--mono);font-size:.62rem;color:var(--silver);letter-spacing:.08em}
.pdf-title{flex:1;font-size:.88rem;color:var(--text);font-family:var(--ser);overflow:hidden;text-overflow:ellipsis;white-space:nowrap;min-width:0}
.pdf-page-info{font-family:var(--mono);font-size:.68rem;color:var(--gold2);letter-spacing:.08em;white-space:nowrap}
.pdf-embed-actions{display:flex;gap:6px;align-items:center;flex-wrap:wrap}
.pdf-btn{font-family:var(--mono);font-size:.65rem;letter-spacing:.08em;color:var(--gold);border:1px solid var(--border2);background:rgba(201,168,76,.06);border-radius:999px;padding:5px 13px;cursor:pointer;text-decoration:none;transition:all var(--trans);white-space:nowrap}
.pdf-btn:hover{background:rgba(201,168,76,.15);color:var(--gold3)}
.pdf-embed-wrap{width:100%;background:#1a1a1a;min-height:500px;position:relative}
.pdf-embed-wrap.fullscreen{position:fixed;inset:0;z-index:500;background:#1a1a1a;overflow-y:auto}
.pdf-canvas-wrap{display:flex;justify-content:center;align-items:flex-start;padding:24px 16px;min-height:500px;position:relative}
.pdf-canvas{max-width:100%;box-shadow:0 8px 40px rgba(0,0,0,.6);border-radius:4px;display:block}
.pdf-loading{position:absolute;inset:0;display:flex;align-items:center;justify-content:center;font-family:var(--mono);font-size:.8rem;color:var(--silver);letter-spacing:.1em}
.card-label{font-family:var(--mono);font-size:.65rem;letter-spacing:.18em;text-transform:uppercase;color:var(--gold);margin-bottom:10px}
.muted{color:var(--text2);font-style:italic;font-size:.88rem}
.card-word,.card-table,.card-json,.card-other{background:var(--bg2);border:1px solid var(--border);border-radius:var(--r);padding:28px 32px;transition:box-shadow var(--trans)}
.card-title{font-size:1.15rem;font-weight:400;color:var(--text);margin-bottom:18px}
.doc-h1{font-size:1rem;color:var(--gold2);margin:14px 0 6px;font-family:var(--ser)}
.doc-h2{font-size:.9rem;color:var(--text);margin:10px 0 5px;font-family:var(--ser)}
.doc-p{font-size:.88rem;color:var(--text2);margin-bottom:8px;font-family:var(--sans)}
.table-wrap{overflow-x:auto}
table{width:100%;border-collapse:collapse;font-size:.82rem;font-family:var(--sans)}
th{padding:10px 14px;text-align:left;font-family:var(--mono);font-size:.6rem;letter-spacing:.12em;text-transform:uppercase;color:var(--gold);border-bottom:1px solid var(--border2);white-space:nowrap}
td{padding:9px 14px;color:var(--text2);border-bottom:1px solid var(--border);white-space:nowrap}
tr:last-child td{border-bottom:none}
tr:hover td{color:var(--text);background:rgba(255,255,255,.025)}
.table-note{font-family:var(--mono);font-size:.62rem;color:var(--silver);margin-top:12px;letter-spacing:.08em}
.json-pre{font-family:var(--mono);font-size:.75rem;color:var(--text2);white-space:pre-wrap;word-break:break-all;background:var(--bg3);padding:20px;border-radius:8px;border:1px solid var(--border);max-height:400px;overflow-y:auto}
.meta{font-family:var(--sans);font-size:.82rem;color:var(--text2);margin-bottom:5px}
footer{position:relative;z-index:1;padding:32px 64px;background:var(--bg2);border-top:1px solid var(--border);display:flex;align-items:center;justify-content:space-between;gap:24px;flex-wrap:wrap;font-family:var(--mono);font-size:.68rem;letter-spacing:.1em;color:var(--silver)}
footer .logo{color:var(--gold);font-family:var(--ser);font-size:.9rem}
footer a{color:var(--text2);text-decoration:none;transition:color var(--trans)}
footer a:hover{color:var(--gold)}
#lightbox{display:none;position:fixed;inset:0;z-index:1000;background:rgba(0,0,0,.92);backdrop-filter:blur(10px);flex-direction:column;align-items:center;justify-content:center}
#lightbox.open{display:flex}
#lightbox-img{max-width:92vw;max-height:82vh;object-fit:contain;border-radius:6px;box-shadow:0 20px 80px rgba(0,0,0,.8)}
#lightbox-caption{margin-top:16px;font-family:var(--mono);font-size:.7rem;letter-spacing:.1em;color:rgba(255,255,255,.55);text-align:center;max-width:600px}
#lightbox-close,#lightbox-prev,#lightbox-next{position:fixed;background:rgba(255,255,255,.1);border:1px solid rgba(255,255,255,.2);color:#fff;border-radius:50%;width:44px;height:44px;cursor:pointer;display:flex;align-items:center;justify-content:center;font-size:1.1rem;transition:background var(--trans)}
#lightbox-close:hover,#lightbox-prev:hover,#lightbox-next:hover{background:rgba(255,255,255,.2)}
#lightbox-close{top:20px;right:20px}
#lightbox-prev{left:20px;top:50%;transform:translateY(-50%)}
#lightbox-next{right:20px;top:50%;transform:translateY(-50%)}
@keyframes fadeUp{from{opacity:0;transform:translateY(24px)}to{opacity:1;transform:translateY(0)}}
@keyframes fadeIn{from{opacity:0}to{opacity:1}}
@keyframes pageFadeIn{from{opacity:0;transform:translateY(18px)}to{opacity:1;transform:translateY(0)}}
.proj-card:nth-child(1){animation-delay:.03s}.proj-card:nth-child(2){animation-delay:.06s}
.proj-card:nth-child(3){animation-delay:.09s}.proj-card:nth-child(4){animation-delay:.12s}
.proj-card:nth-child(5){animation-delay:.15s}.proj-card:nth-child(6){animation-delay:.18s}
.proj-card:nth-child(7){animation-delay:.21s}.proj-card:nth-child(8){animation-delay:.24s}
@media(max-width:768px){.site-nav{padding:0 16px}.site-header,.page-header{padding:36px 20px 28px}.home-grid{padding:28px 16px 50px;grid-template-columns:1fr}.tab-bar,.tab-content{padding-left:16px;padding-right:16px}footer{padding:24px 20px;flex-direction:column;align-items:flex-start}.gallery{grid-template-columns:repeat(auto-fill,minmax(150px,1fr))}}
"""


def _html_foot(now: str) -> str:
    return f"""
<!-- LIGHTBOX -->
<div id="lightbox" role="dialog" aria-modal="true">
  <button id="lightbox-close" onclick="closeLightbox()" aria-label="Kapat">✕</button>
  <button id="lightbox-prev" onclick="lightboxStep(-1)" aria-label="Önceki">‹</button>
  <img id="lightbox-img" src="" alt=""/>
  <div id="lightbox-caption"></div>
  <button id="lightbox-next" onclick="lightboxStep(1)" aria-label="Sonraki">›</button>
</div>

<!-- FOOTER -->
<footer>
  <span class="logo">Erdoğan Işık Özel · <em>CB VIP Dikmen</em></span>
  <div style="display:flex;gap:24px;align-items:center;flex-wrap:wrap">
    <a href="/">Ana Sayfa</a>
    <a href="/ilanlar">İlanlar</a>
    <a href="https://wa.me/905324514008" target="_blank">WhatsApp</a>
  </div>
  <span>Nexa CRM · {now}</span>
</footer>

<script>
function pdfPrev(uid) {{ const s=window['_pdfState_'+uid]; if(s) s.prev(); }}
function pdfNext(uid) {{ const s=window['_pdfState_'+uid]; if(s) s.next(); }}
function togglePdfFull(uid) {{
  const wrap=document.getElementById('pdfwrap-'+uid);
  if(!wrap) return;
  wrap.classList.toggle('fullscreen');
  document.body.style.overflow=wrap.classList.contains('fullscreen')?'hidden':'';
}}
function openProjectAnim(cardEl,slug) {{
  const rect=cardEl.getBoundingClientRect();
  const cx=rect.left+rect.width/2,cy=rect.top+rect.height/2;
  const ov=document.createElement('div');
  ov.style.cssText='position:fixed;left:'+cx+'px;top:'+cy+'px;width:4px;height:4px;border-radius:50%;background:var(--bg2);transform:translate(-50%,-50%) scale(0);transition:transform .55s cubic-bezier(.4,0,.2,1),opacity .15s ease .45s;z-index:900;pointer-events:none;';
  document.body.appendChild(ov);
  const maxR=Math.sqrt(Math.pow(Math.max(cx,window.innerWidth-cx)*2,2)+Math.pow(Math.max(cy,window.innerHeight-cy)*2,2));
  requestAnimationFrame(()=>{{ov.style.transform='translate(-50%,-50%) scale('+maxR+')';ov.style.opacity='1';}});
  setTimeout(()=>{{openProject(slug);ov.style.opacity='0';setTimeout(()=>ov.remove(),200);}},480);
}}
function openProject(slug) {{
  document.getElementById('homepage').style.display='none';
  document.querySelectorAll('.project-page').forEach(p=>p.style.display='none');
  const pg=document.getElementById('page-'+slug);
  if(pg){{pg.style.display='block';pg.style.animation='pageFadeIn .35s ease both';window.scrollTo(0,0);}}
}}
function closeProject() {{
  document.querySelectorAll('.project-page').forEach(p=>p.style.display='none');
  document.getElementById('homepage').style.display='block';
  window.scrollTo(0,0);
}}
function switchTab(slug,key) {{
  const page=document.getElementById('page-'+slug);
  if(!page) return;
  page.querySelectorAll('.tab-panel').forEach(p=>p.style.display='none');
  page.querySelectorAll('.tab-btn').forEach(b=>b.classList.remove('active'));
  const panel=document.getElementById('panel-'+slug+'-'+key);
  if(panel) panel.style.display='';
  event.target.classList.add('active');
}}
document.addEventListener('keydown',e=>{{
  if(e.key==='Escape'){{
    if(document.querySelectorAll('.pdf-embed-wrap.fullscreen').length){{
      document.querySelectorAll('.pdf-embed-wrap.fullscreen').forEach(w=>w.classList.remove('fullscreen'));
      document.body.style.overflow='';
    }} else if(document.getElementById('lightbox').classList.contains('open')) closeLightbox();
    else closeProject();
  }}
}});
let _lbItems=[],_lbIdx=0;
function openLightbox(el) {{
  const gallery=el.closest('.gallery');
  _lbItems=Array.from(gallery.querySelectorAll('.gal-item'));
  _lbIdx=_lbItems.indexOf(el);
  showLightboxItem();
  document.getElementById('lightbox').classList.add('open');
  document.body.style.overflow='hidden';
}}
function showLightboxItem() {{
  const item=_lbItems[_lbIdx];
  const img=item.querySelector('img');
  const lb=document.getElementById('lightbox-img');
  lb.src=img.dataset.large||img.src;
  lb.alt=img.alt;
  document.getElementById('lightbox-caption').textContent=img.alt+(_lbItems.length>1?'  ('+(_lbIdx+1)+' / '+_lbItems.length+')':'');
}}
function lightboxStep(dir) {{
  _lbIdx=(_lbIdx+dir+_lbItems.length)%_lbItems.length;
  showLightboxItem();
}}
function closeLightbox() {{
  document.getElementById('lightbox').classList.remove('open');
  document.body.style.overflow='';
}}
document.getElementById('lightbox').addEventListener('click',function(e){{if(e.target===this)closeLightbox();}});
document.addEventListener('keydown',e=>{{
  if(!document.getElementById('lightbox').classList.contains('open')) return;
  if(e.key==='ArrowRight') lightboxStep(1);
  if(e.key==='ArrowLeft')  lightboxStep(-1);
}});
// ── Drive Refresh (Admin) ────────────────────────────────
function driveRefresh(btn) {{
  if(!btn) return;
  btn.classList.add('spinning');
  btn.disabled=true;
  fetch('/api/drive/refresh', {{method:'POST'}})
    .then(r=>r.json())
    .then(d=>{{
      btn.classList.remove('spinning');
      btn.disabled=false;
      if(d.ok) {{
        btn.textContent='✓ Yenilendi';
        setTimeout(()=>location.reload(), 1200);
      }} else {{
        btn.textContent='Hata: '+d.error;
        setTimeout(()=>{{btn.textContent='↻ Yenile'; }}, 3000);
      }}
    }})
    .catch(()=>{{ btn.classList.remove('spinning'); btn.disabled=false; }});
}}
</script>
</body>
</html>
"""


# ════════════════════════════════════════════════════════════
#  10. ANA FONKSİYON
# ════════════════════════════════════════════════════════════
def main():
    print("─" * 52)
    print("  drive_to_html.py — Lüks HTML Sunum Üretici")
    print("─" * 52)

    print("\n[1/4] Google Drive'a bağlanılıyor…")
    Path(ASSETS_DIR).mkdir(exist_ok=True)
    service = get_service()
    print("  ✓ Bağlantı başarılı")

    print(f"\n[2/4] Klasör taranıyor (id: {FOLDER_ID})…")
    files = list_files(service, FOLDER_ID)
    if not files:
        print("  ⚠  Klasörde dosya bulunamadı.")
        return
    print(f"  ✓ {len(files)} dosya bulundu")

    # ── Değişiklik kontrolü ──────────────────────────────────
    current_hash = compute_manifest(files)
    prev_hash    = load_manifest()
    html_exists  = Path(OUTPUT_FILE).exists()

    if current_hash == prev_hash and html_exists:
        print(f"\n  ✅ Drive'da değişiklik yok — {OUTPUT_FILE} zaten güncel.")
        print(f"     (Yeniden üretmek için '{MANIFEST_FILE}' dosyasını sil.)")
        print("─" * 52)
        return
    elif prev_hash and current_hash != prev_hash:
        print(f"  🔄 Değişiklik tespit edildi — HTML yeniden üretiliyor…")
    else:
        print(f"  🆕 İlk çalıştırma — HTML üretiliyor…")

    print("\n[3/4] Dosyalar işleniyor…")
    stats = cache_stats()
    if stats["files"] > 0:
        print(f"  ℹ  Önbellek: {stats['files']} dosya, {stats['size_mb']:.1f} MB  ({CACHE_DIR}/)")
    else:
        print(f"  ℹ  Önbellek boş — ilk çalıştırmada tüm dosyalar indirilecek")
    _cache_hits = [0]
    _cache_miss = [0]

    # ── Proje gruplarını oluştur ──────────────────────────────
    from collections import OrderedDict
    projects = OrderedDict()   # project_name → dict of lists

    def get_project(fp: str) -> str:
        parts = fp.split(" / ")
        return parts[0] if parts else "DİĞER"

    for f in files:
        pname = get_project(f.get("folder_path", ""))
        if pname not in projects:
            projects[pname] = {
                "images":  [],   # (thumb_src, large_src, name, sub_folder)
                "pdfs":    [],   # (html_card,)
                "tables":  [],   # (html_card,)
                "words":   [],   # (html_card,)
                "others":  [],   # (html_card,)
                "cover":   "",   # first image thumb
                "video":   "",   # base64 data URI for cover video (mp4)
                "total":   0,
            }
        projects[pname]["total"] += 1

    total = len(files)
    for i, f in enumerate(files, 1):
        name = f["name"]
        mime = f.get("mimeType", "")
        fid  = f["id"]
        tkey = file_type_key(name, mime)
        fp   = f.get("folder_path", "")
        pname = get_project(fp)
        sub   = " / ".join(fp.split(" / ")[1:]) if " / " in fp else ""
        print(f"  [{i:03}/{total:03}] {name[:48]:<48} [{tkey}]")

        try:
            _prev_files = cache_stats()["files"]
            data, real_mime = download_cached(service, fid, mime, f.get("modifiedTime",""))
            if cache_stats()["files"] > _prev_files: _cache_miss[0] += 1
            else: _cache_hits[0] += 1
            tkey2 = file_type_key(name, real_mime) if real_mime != mime else tkey

            # ── CONTENT FILTER ──────────────────────────────────
            name_up = name.upper()
            # Unique asset ID for this file
            asset_uid = hashlib.md5((fid + f.get("modifiedTime","")).encode()).hexdigest()[:12]

            if tkey2 == "other":
                is_video = (real_mime.startswith("video/") or name.lower().endswith(".mp4"))
                is_tanitim = "TANITIM" in name_up
                if is_video and is_tanitim and not projects[pname]["video"]:
                    vpath = save_video(data, asset_uid)
                    if vpath:
                        projects[pname]["video"] = vpath
                        print(f"    ✓ Video kaydedildi: {name}")
                del data  # RAM'i serbest bırak
                continue
            if tkey2 in ("word", "table", "json"):
                del data
                continue
            if tkey2 == "pdf":
                if not any(k in name_up for k in (
                    "SUNUM", "FİYAT", "FIYAT", "ÖDEME", "ODEME",
                    "LİSTE", "LISTE", "TANITIM", "KATALOG"
                )):
                    del data
                    continue

            if tkey2 == "image":
                thumb = make_thumb(data, real_mime, asset_uid)
                large = make_large(data, real_mime, asset_uid)
                del data  # RAM'i hemen serbest bırak
                projects[pname]["images"].append((thumb, large, html.escape(name), html.escape(sub), html.escape(fp)))
                is_exterior = "DIŞ CEPHE" in fp.upper() or "DIS CEPHE" in fp.upper()
                if not projects[pname]["cover"] and thumb:
                    projects[pname]["cover"] = thumb
                    projects[pname]["cover_is_exterior"] = is_exterior
                elif is_exterior and not projects[pname].get("cover_is_exterior") and thumb:
                    projects[pname]["cover"] = thumb
                    projects[pname]["cover_is_exterior"] = True

            elif tkey2 == "pdf":
                lbl = "Slides" if mime == "application/vnd.google-apps.presentation" else "PDF"
                card = process_pdf(data, name, lbl)
                del data  # RAM'i hemen serbest bırak
                projects[pname]["pdfs"].append(card)

        except Exception as e:
            print(f"    ⚠  İşlenemedi: {e}")
            try: del data
            except: pass

    # ── HTML stream yaz ──────────────────────────────────────
    print(f"\n  Önbellek özeti: {_cache_hits[0]} hit (atlandı), {_cache_miss[0]} miss (indirildi)")
    print(f"[4/4] HTML dosyası yazılıyor → {OUTPUT_FILE}")
    now    = datetime.now().strftime("%d.%m.%Y %H:%M")
    n_proj = len(projects)
    _words = PROJE_ADI.split()
    proje_adi_h1 = " ".join(_words[:-1]) + (f" <em>{_words[-1]}</em>" if len(_words) > 1 else _words[0])

    # ── CSS/JS şablonunu ayrı dosyadan değil, fonksiyonlardan çek ──────────
    HTML_HEAD   = _html_head(PROJE_ADI)
    HTML_CSS    = _html_css()
    HTML_FOOT   = _html_foot(now)

    with open(OUTPUT_FILE, "w", encoding="utf-8") as fh:

        # 1. HEAD + CSS
        fh.write(HTML_HEAD)
        fh.write(HTML_CSS)
        fh.write("</style>\n</head>\n<body>\n\n")

        # 2. Nav bar
        fh.write("""<nav class="site-nav">
  <a href="/" class="nav-logo">Erdoğan <em>Işık</em> Özel</a>
  <div class="nav-links">
    <a href="/">Ana Sayfa</a>
    <a href="/ilanlar"><i class="fas fa-building" style="margin-right:4px;font-size:10px;"></i>İlanlar</a>
    <a href="/portfoy" class="active"><i class="fas fa-layer-group" style="margin-right:4px;font-size:10px;"></i>Portföy</a>
    <a href="/#contact">İletişim</a>
    <a href="https://wa.me/905324514008" target="_blank" class="nav-wa">
      <i class="fab fa-whatsapp" style="margin-right:4px;"></i>WhatsApp
    </a>
  </div>
</nav>\n\n""")

        # 3. Site header
        fh.write(f"""<header class="site-header" id="site-header">
  <div class="header-inner">
    <p class="header-eyebrow">CB VIP Dikmen · Proje Portföyü · Drive Arşivi</p>
    <h1>{proje_adi_h1}</h1>
    <p class="header-sub">{html.escape(PROJE_ALT_BASLIK)}</p>
    <div class="header-meta">
      <div class="hm-pill">Proje: <span>{n_proj}</span></div>
      <div class="hm-pill">Toplam Dosya: <span>{total}</span></div>
      <div class="hm-pill">Oluşturulma: <span>{now}</span></div>
      <button class="header-refresh" onclick="driveRefresh(this)">
        <i class="fas fa-sync-alt"></i> Yenile
      </button>
    </div>
  </div>
</header>\n\n""")

        # 3. Ana sayfa kartları
        fh.write('<div id="homepage">\n  <div class="home-grid">\n')
        for pname, pdata in projects.items():
            slug    = slugify(pname)
            cover   = pdata["cover"] or ""
            n_img   = len(pdata["images"])
            n_pdf   = len(pdata["pdfs"])
            total_p = pdata["total"]
            video_src = pdata.get("video", "")
            has_video = bool(video_src)

            if has_video:
                cover_html = (f'<video class="proj-cover proj-cover-video" '
                              f'src="{video_src}" autoplay muted loop playsinline '
                              f'poster="{cover}" preload="metadata"></video>'
                              f'<div class="proj-video-badge">▶ Video</div>')
            elif cover:
                cover_html = f'<img src="{cover}" class="proj-cover" alt="{html.escape(pname)}" loading="lazy"/>'
            else:
                cover_html = '<div class="proj-cover proj-cover-placeholder"><span>◈</span></div>'

            badges = ""
            if has_video: badges += '<span class="badge badge-video">▶ Video</span>'
            if n_img:     badges += f'<span class="badge">🖼 {n_img} görsel</span>'
            if n_pdf:     badges += f'<span class="badge">📄 {n_pdf} PDF</span>'
            hint = '▶ Tanıtımı İzle →' if has_video else 'Görüntüle →'

            fh.write(f"""<div class="proj-card" onclick="openProjectAnim(this,'{slug}')" tabindex="0" role="button"
     onkeydown="if(event.key==='Enter')openProjectAnim(this,'{slug}')"
     aria-label="{html.escape(pname)}">
  <div class="proj-cover-wrap">{cover_html}<div class="proj-overlay"><span class="proj-open-hint">{hint}</span></div></div>
  <div class="proj-info">
    <div class="proj-eyebrow">Proje</div>
    <h2 class="proj-name">{html.escape(pname)}</h2>
    <div class="proj-badges">{badges}</div>
    <div class="proj-total">{total_p} dosya</div>
  </div>
</div>\n""")

        fh.write('  </div>\n</div>\n\n<!-- PROJECT PAGES -->\n')

        # 4. Proje detay sayfaları
        for pname, pdata in projects.items():
            slug    = slugify(pname)
            n_img   = len(pdata["images"])
            n_pdf   = len(pdata["pdfs"])
            total_p = pdata["total"]
            has_video = bool(pdata.get("video"))

            badges = ""
            if has_video: badges += '<span class="badge badge-video">▶ Video</span>'
            if n_img:     badges += f'<span class="badge">🖼 {n_img} görsel</span>'
            if n_pdf:     badges += f'<span class="badge">📄 {n_pdf} PDF</span>'

            # Tab bar
            tabs_html = ""
            if pdata["images"]: tabs_html += f'<button class="tab-btn active" onclick="switchTab(\'{slug}\',\'images\')">🖼 Görseller <span class="tab-count">{n_img}</span></button>'
            if pdata["pdfs"]:
                active2 = " active" if not pdata["images"] else ""
                tabs_html += f'<button class="tab-btn{active2}" onclick="switchTab(\'{slug}\',\'pdfs\')">📄 Sunum &amp; Fiyat <span class="tab-count">{n_pdf}</span></button>'

            fh.write(f"""<section class="project-page" id="page-{slug}" style="display:none">
  <div class="page-header">
    <button class="back-btn" onclick="closeProject()">← Tüm Projeler</button>
    <div class="page-header-inner">
      <div class="page-eyebrow">Proje Detayı</div>
      <h1 class="page-title">{html.escape(pname)}</h1>
      <div class="page-stats">{badges}<span class="badge badge-total">{total_p} dosya</span></div>
    </div>
  </div>
  <div class="tab-bar" id="tabs-{slug}">{tabs_html}</div>
  <div class="tab-content">\n""")

            # Görseller paneli
            if pdata["images"]:
                fh.write(f'<div class="tab-panel" id="panel-{slug}-images">\n<div class="gallery">\n')
                for item in pdata["images"]:
                    thumb, large, iname, isub = item[0], item[1], item[2], item[3]
                    if not thumb:
                        continue  # başarısız görseli atla
                    sub_tag = f'<span class="gal-sub">{isub}</span>' if isub else ""
                    fh.write(f'<div class="gal-item" onclick="openLightbox(this)">'
                             f'<img src="{thumb}" data-large="{large}" alt="{iname}" loading="lazy"/>'
                             f'<div class="gal-caption">{iname}{sub_tag}</div></div>\n')
                fh.write('</div>\n</div>\n')  # gallery + tab-panel

            # PDF paneli
            if pdata["pdfs"]:
                hidden = ' style="display:none"' if pdata["images"] else ""
                fh.write(f'<div class="tab-panel" id="panel-{slug}-pdfs"{hidden}>\n<div class="pdf-stack">\n')
                for card in pdata["pdfs"]:
                    fh.write(card + "\n")
                fh.write('</div>\n</div>\n')  # pdf-stack + tab-panel

            fh.write('  </div>\n</section>\n\n')  # tab-content + section

        # 5. Lightbox + footer + JS
        fh.write(HTML_FOOT)

    size_kb = Path(OUTPUT_FILE).stat().st_size / 1024
    asset_files = len(list(Path(ASSETS_DIR).iterdir())) if Path(ASSETS_DIR).exists() else 0
    save_manifest(current_hash)
    print(f"  ✓ {OUTPUT_FILE} oluşturuldu ({size_kb:.0f} KB)")
    print(f"  ✓ {ASSETS_DIR}/ klasörü: {asset_files} dosya")
    print(f"\n  ⚠  HTML ve '{ASSETS_DIR}/' klasörünü birlikte tut!")
    print(f"\n  Tarayıcıda aç:")
    print(f"    python -m http.server 8080  →  http://localhost:8080/{OUTPUT_FILE}")
    print("─" * 52)


if __name__ == "__main__":
    main()
